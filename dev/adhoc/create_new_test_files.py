#!/usr/bin/env python3
"""
Create new test DOCX files - one unencrypted and one encrypted
"""

from docx import Document
import msoffcrypto
import io
from pathlib import Path
import random

def create_random_content():
    """Generate random content for the documents"""
    topics = [
        "Market Analysis", "Project Proposal", "Technical Specifications", 
        "Financial Report", "Strategic Planning", "Research Findings",
        "Implementation Guide", "Performance Review", "Risk Assessment"
    ]
    
    companies = [
        "TechCorp Solutions", "DataFlow Industries", "Innovation Labs",
        "Strategic Ventures", "Global Systems", "NextGen Technologies"
    ]
    
    departments = [
        "Engineering", "Marketing", "Finance", "Operations", 
        "Research & Development", "Quality Assurance"
    ]
    
    return {
        "title": random.choice(topics),
        "company": random.choice(companies),
        "department": random.choice(departments),
        "content": [
            f"This document presents a comprehensive analysis of {random.choice(['current market trends', 'operational efficiency', 'strategic initiatives', 'technological developments'])}.",
            f"Our {random.choice(departments).lower()} team has identified {random.randint(3, 8)} key areas for improvement.",
            f"The proposed timeline spans {random.randint(6, 18)} months with an estimated budget of ${random.randint(50, 500)}K.",
            f"Key stakeholders include {random.choice(['senior management', 'project teams', 'external partners', 'regulatory bodies'])}.",
            f"Expected outcomes include {random.choice(['increased revenue', 'improved efficiency', 'enhanced security', 'better compliance'])} by {random.randint(15, 40)}%."
        ]
    }

def create_unencrypted_docx():
    """Create an unencrypted DOCX file with random content"""
    print("Creating unencrypted DOCX file...")
    
    content = create_random_content()
    doc = Document()
    
    # Add title
    title = doc.add_heading(f'{content["title"]} - {content["company"]}', 0)
    
    # Add metadata paragraph
    doc.add_paragraph(f'Department: {content["department"]}')
    doc.add_paragraph(f'Document Type: Internal Report')
    doc.add_paragraph(f'Classification: Public')
    
    # Add main content
    doc.add_heading('Executive Summary', level=1)
    for paragraph in content["content"]:
        doc.add_paragraph(paragraph)
    
    # Add a simple table
    doc.add_heading('Project Timeline', level=1)
    table = doc.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Phase'
    hdr_cells[1].text = 'Duration'
    hdr_cells[2].text = 'Status'
    
    phases = ['Planning', 'Implementation', 'Testing', 'Deployment']
    for i, phase in enumerate(phases):
        row_cells = table.add_row().cells
        row_cells[0].text = phase
        row_cells[1].text = f'{random.randint(2, 8)} weeks'
        row_cells[2].text = random.choice(['Pending', 'In Progress', 'Complete'])
    
    # Save the document
    output_path = Path(__file__).parent / 'unencrypted_document.docx'
    doc.save(str(output_path))
    print(f"Created unencrypted file: {output_path}")
    return output_path

def create_encrypted_docx():
    """Create an encrypted DOCX file with random content"""
    print("Creating encrypted DOCX file...")
    
    content = create_random_content()
    doc = Document()
    
    # Add title
    title = doc.add_heading(f'CONFIDENTIAL: {content["title"]}', 0)
    
    # Add metadata paragraph
    doc.add_paragraph(f'Company: {content["company"]}')
    doc.add_paragraph(f'Department: {content["department"]}')
    doc.add_paragraph(f'Document Type: Confidential Report')
    doc.add_paragraph(f'Classification: RESTRICTED ACCESS')
    
    # Add warning
    warning = doc.add_paragraph()
    warning.add_run('CONFIDENTIAL INFORMATION').bold = True
    doc.add_paragraph('This document contains sensitive information and is protected by password encryption.')
    
    # Add main content
    doc.add_heading('Confidential Analysis', level=1)
    for paragraph in content["content"]:
        doc.add_paragraph(paragraph)
    
    # Add sensitive data section
    doc.add_heading('Sensitive Data', level=1)
    doc.add_paragraph(f'Internal Reference ID: {random.randint(100000, 999999)}')
    doc.add_paragraph(f'Security Level: {random.choice(["Level 1", "Level 2", "Level 3"])}')
    doc.add_paragraph(f'Access Code: {random.randint(1000, 9999)}')
    
    # Save to memory first
    temp_output = io.BytesIO()
    doc.save(temp_output)
    temp_output.seek(0)
    
    # Encrypt the document with password "password"
    print("Encrypting document with password 'password'...")
    encrypted_output = io.BytesIO()
    
    office_file = msoffcrypto.OfficeFile(temp_output)
    office_file.encrypt("password", encrypted_output)
    
    # Save encrypted file
    output_path = Path(__file__).parent / 'encrypted_document.docx'
    with open(output_path, 'wb') as f:
        f.write(encrypted_output.getvalue())
    
    print(f"Created encrypted file: {output_path}")
    print("Password: password")
    return output_path

if __name__ == "__main__":
    print("Creating new test DOCX files with random content...\n")
    
    # Create both files
    unencrypted_file = create_unencrypted_docx()
    encrypted_file = create_encrypted_docx()
    
    print(f"\nFiles created in {Path(__file__).parent}:")
    print(f"  Unencrypted: {unencrypted_file.name}")
    print(f"  Encrypted: {encrypted_file.name} (password: 'password')")
    print("\nBoth files contain random business content for testing purposes.")