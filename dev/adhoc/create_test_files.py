#!/usr/bin/env python3
"""
Create test DOCX files for FastPass testing
"""

from docx import Document
import msoffcrypto
import io
from pathlib import Path

def create_unencrypted_file():
    """Create file1_unencrypted.docx with sample content"""
    doc = Document()
    
    # Add title
    title = doc.add_heading('Sample Business Document', 0)
    
    # Add paragraphs
    doc.add_heading('Executive Summary', level=1)
    doc.add_paragraph(
        'This is a sample business document created for testing the FastPass '
        'encryption tool. The document contains various formatting elements '
        'and content to demonstrate the tool\'s capabilities.'
    )
    
    doc.add_heading('Project Details', level=1)
    doc.add_paragraph(
        'Project Name: FastPass Testing Initiative\n'
        'Start Date: January 2024\n'
        'Status: Active\n'
        'Priority: High'
    )
    
    # Add a table
    table = doc.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Task'
    hdr_cells[1].text = 'Owner'
    hdr_cells[2].text = 'Status'
    
    # Add data rows
    row_cells = table.add_row().cells
    row_cells[0].text = 'File Encryption Testing'
    row_cells[1].text = 'Security Team'
    row_cells[2].text = 'In Progress'
    
    row_cells = table.add_row().cells
    row_cells[0].text = 'Documentation Review'
    row_cells[1].text = 'Tech Writers'
    row_cells[2].text = 'Complete'
    
    # Add bullet points
    doc.add_heading('Key Features', level=1)
    doc.add_paragraph('Universal file encryption support', style='List Bullet')
    doc.add_paragraph('Password-based protection', style='List Bullet')
    doc.add_paragraph('Multiple file format compatibility', style='List Bullet')
    doc.add_paragraph('Secure temporary file handling', style='List Bullet')
    
    # Save the document
    output_path = Path(__file__).parent / 'file1_unencrypted.docx'
    doc.save(str(output_path))
    print(f"Created: {output_path}")
    return output_path

def create_encrypted_file():
    """Create file2_encrypted.docx with different content and password protection"""
    doc = Document()
    
    # Add title
    title = doc.add_heading('Confidential Security Report', 0)
    
    # Add content
    doc.add_heading('Security Assessment Results', level=1)
    doc.add_paragraph(
        'CONFIDENTIAL: This document contains sensitive security information '
        'and should be protected at all times. Access is restricted to '
        'authorized personnel only.'
    )
    
    doc.add_heading('Risk Analysis', level=1)
    doc.add_paragraph(
        'The following risks have been identified:\n'
        '• High: Unauthorized access to encrypted files\n'
        '• Medium: Password sharing between users\n'
        '• Low: Temporary file exposure during processing'
    )
    
    doc.add_heading('Recommendations', level=1)
    doc.add_paragraph(
        '1. Implement strong password policies\n'
        '2. Regular security audits\n'
        '3. Employee training on data protection\n'
        '4. Automated encryption for sensitive documents'
    )
    
    doc.add_paragraph(
        'This report was generated as part of the FastPass security testing '
        'initiative to validate encryption and decryption capabilities.'
    )
    
    # Save to memory first
    temp_output = io.BytesIO()
    doc.save(temp_output)
    temp_output.seek(0)
    
    # Encrypt the document
    encrypted_output = io.BytesIO()
    
    office_file = msoffcrypto.OfficeFile(temp_output)
    office_file.encrypt("testpassword", encrypted_output)
    
    # Save encrypted file
    output_path = Path(__file__).parent / 'file2_encrypted.docx'
    with open(output_path, 'wb') as f:
        f.write(encrypted_output.getvalue())
    
    print(f"Created encrypted file: {output_path}")
    print("Password: testpassword")
    return output_path

if __name__ == "__main__":
    print("Creating test DOCX files...")
    file1 = create_unencrypted_file()
    file2 = create_encrypted_file()
    print(f"\nFiles created:")
    print(f"  Unencrypted: {file1}")
    print(f"  Encrypted: {file2} (password: testpassword)")